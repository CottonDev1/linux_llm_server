"""
Document Processor Service

Extracts text and metadata from various document types (PDF, DOCX, XLSX, TXT)
for storage in the MongoDB RAG system.

Follows best practices from 2025 research:
- Docling/PyMuPDF for PDF with table extraction
- python-docx for Word documents
- openpyxl for Excel spreadsheets
- Metadata-enhanced embeddings for better retrieval
"""

import os
import hashlib
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import asyncio


@dataclass
class DocumentMetadata:
    """Extracted document metadata"""
    title: str = ""
    author: str = ""
    created_date: str = ""
    page_count: int = 0
    word_count: int = 0
    has_tables: bool = False
    has_images: bool = False
    table_count: int = 0
    sheet_names: List[str] = field(default_factory=list)
    language: str = "en"
    file_hash: str = ""


@dataclass
class ProcessedDocument:
    """Result of document processing"""
    content: str
    content_type: str
    metadata: DocumentMetadata
    file_name: str
    file_size: int
    success: bool = True
    error: Optional[str] = None


class DocumentProcessor:
    """
    Extract text and metadata from various document types.
    Integrates with MongoDB RAG system for document storage.
    """

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.xlsx': 'xlsx',
        '.xls': 'xls',
        '.txt': 'txt',
        '.md': 'markdown',
        '.csv': 'csv',
        '.json': 'json'
    }

    MIME_TYPE_MAP = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
        'application/vnd.ms-excel': 'xls',
        'text/plain': 'txt',
        'text/markdown': 'markdown',
        'text/csv': 'csv',
        'application/json': 'json'
    }

    # Security limits
    DEFAULT_MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
    DEFAULT_ALLOWED_DIRS: List[str] = []  # Empty = allow all (for backward compat)

    def __init__(
        self,
        max_file_size: int = DEFAULT_MAX_FILE_SIZE,
        allowed_dirs: Optional[List[str]] = None,
    ):
        """
        Initialize document processor with security settings.

        Args:
            max_file_size: Maximum file size in bytes (default: 100MB)
            allowed_dirs: List of allowed directories. If empty, all paths allowed.
        """
        self._max_file_size = max_file_size
        self._allowed_dirs = [Path(d).resolve() for d in (allowed_dirs or [])]
        self._check_dependencies()

    def _check_dependencies(self):
        """Check which optional dependencies are available"""
        self.has_pymupdf = False
        self.has_docx = False
        self.has_openpyxl = False
        self.has_magic = False

        try:
            import fitz  # PyMuPDF
            self.has_pymupdf = True
        except ImportError:
            print("Warning: PyMuPDF not installed. PDF processing will be limited.")

        try:
            from docx import Document
            self.has_docx = True
        except ImportError:
            print("Warning: python-docx not installed. DOCX processing unavailable.")

        try:
            import openpyxl
            self.has_openpyxl = True
        except ImportError:
            print("Warning: openpyxl not installed. XLSX processing unavailable.")

        try:
            import magic
            self.has_magic = True
        except ImportError:
            pass  # Will use extension-based detection

    def _validate_path(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate file path for security (path traversal protection).

        Args:
            file_path: Path to validate

        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            resolved = Path(file_path).resolve()

            # Check for path traversal attempts
            if ".." in str(file_path):
                return False, "Path traversal detected: '..' not allowed in path"

            # If allowed_dirs is configured, check path is within allowed directories
            if self._allowed_dirs:
                is_allowed = any(
                    resolved == allowed_dir or allowed_dir in resolved.parents
                    for allowed_dir in self._allowed_dirs
                )
                if not is_allowed:
                    return False, f"Path not in allowed directories: {resolved}"

            # Ensure it's a file (not directory)
            if resolved.exists() and resolved.is_dir():
                return False, "Path is a directory, not a file"

            return True, None

        except Exception as e:
            return False, f"Path validation error: {str(e)}"

    def _validate_file_size(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate file size is within limits.

        Args:
            file_path: Path to file

        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            file_size = os.path.getsize(file_path)
            if file_size > self._max_file_size:
                max_mb = self._max_file_size / (1024 * 1024)
                actual_mb = file_size / (1024 * 1024)
                return False, f"File size ({actual_mb:.1f}MB) exceeds maximum ({max_mb:.1f}MB)"
            return True, None
        except OSError as e:
            return False, f"Cannot read file size: {str(e)}"

    def detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect file type from extension or magic bytes"""
        # Try extension first
        ext = Path(file_path).suffix.lower()
        if ext in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[ext]

        # Try magic if available
        if self.has_magic:
            try:
                import magic
                mime = magic.Magic(mime=True)
                mime_type = mime.from_file(file_path)
                return self.MIME_TYPE_MAP.get(mime_type)
            except Exception:
                pass

        return None

    def compute_file_hash(self, file_path: str) -> str:
        """Compute SHA256 hash of file for deduplication"""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                sha256.update(chunk)
        return sha256.hexdigest()

    async def process_file(
        self,
        file_path: str,
        original_filename: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process a document file and extract text + metadata.

        Args:
            file_path: Path to the file on disk
            original_filename: Original filename (for uploaded files)

        Returns:
            ProcessedDocument with content, metadata, and processing status
        """
        file_name = original_filename or os.path.basename(file_path)

        # Security validation: check path traversal
        path_valid, path_error = self._validate_path(file_path)
        if not path_valid:
            return ProcessedDocument(
                content="",
                content_type="unknown",
                metadata=DocumentMetadata(),
                file_name=file_name,
                file_size=0,
                success=False,
                error=f"Security error: {path_error}"
            )

        # Check file exists
        if not os.path.exists(file_path):
            return ProcessedDocument(
                content="",
                content_type="unknown",
                metadata=DocumentMetadata(),
                file_name=file_name,
                file_size=0,
                success=False,
                error="File not found"
            )

        # Security validation: check file size
        size_valid, size_error = self._validate_file_size(file_path)
        if not size_valid:
            return ProcessedDocument(
                content="",
                content_type="unknown",
                metadata=DocumentMetadata(),
                file_name=file_name,
                file_size=os.path.getsize(file_path),
                success=False,
                error=f"Security error: {size_error}"
            )

        file_size = os.path.getsize(file_path)

        # Detect file type
        content_type = self.detect_file_type(file_path)
        if not content_type:
            return ProcessedDocument(
                content="",
                content_type="unknown",
                metadata=DocumentMetadata(),
                file_name=file_name,
                file_size=file_size,
                success=False,
                error=f"Unsupported file type: {Path(file_path).suffix}"
            )

        # Compute file hash for deduplication
        file_hash = self.compute_file_hash(file_path)

        try:
            # Process based on type
            if content_type == 'pdf':
                content, metadata = await self._process_pdf(file_path)
            elif content_type in ('docx', 'doc'):
                content, metadata = await self._process_docx(file_path)
            elif content_type in ('xlsx', 'xls'):
                content, metadata = await self._process_xlsx(file_path)
            elif content_type in ('txt', 'markdown', 'csv', 'json'):
                content, metadata = await self._process_text(file_path)
            else:
                return ProcessedDocument(
                    content="",
                    content_type=content_type,
                    metadata=DocumentMetadata(file_hash=file_hash),
                    file_name=file_name,
                    file_size=file_size,
                    success=False,
                    error=f"No processor available for type: {content_type}"
                )

            metadata.file_hash = file_hash
            metadata.word_count = len(content.split())

            return ProcessedDocument(
                content=content,
                content_type=content_type,
                metadata=metadata,
                file_name=file_name,
                file_size=file_size,
                success=True
            )

        except Exception as e:
            return ProcessedDocument(
                content="",
                content_type=content_type,
                metadata=DocumentMetadata(file_hash=file_hash),
                file_name=file_name,
                file_size=file_size,
                success=False,
                error=str(e)
            )

    async def _process_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from PDF using PyMuPDF"""
        if not self.has_pymupdf:
            raise ImportError("PyMuPDF (fitz) is required for PDF processing. Install with: pip install pymupdf")

        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        metadata = DocumentMetadata()

        try:
            text_parts = []
            metadata.page_count = len(doc)
            metadata.has_tables = False
            metadata.has_images = False

            for page_num, page in enumerate(doc):
                # Extract text
                text = page.get_text()
                text_parts.append(text)

                # Check for images
                if page.get_images():
                    metadata.has_images = True

                # Check for tables (basic detection via line analysis)
                tables = page.find_tables()
                if tables and len(tables.tables) > 0:
                    metadata.has_tables = True
                    metadata.table_count += len(tables.tables)

                    # Extract table content
                    for table in tables:
                        try:
                            table_text = []
                            for row in table.extract():
                                row_text = [str(cell) if cell else '' for cell in row]
                                table_text.append(' | '.join(row_text))
                            text_parts.append('\n' + '\n'.join(table_text))
                        except Exception:
                            pass  # Skip malformed tables

            # Extract document metadata
            doc_metadata = doc.metadata
            if doc_metadata:
                metadata.author = doc_metadata.get('author', '') or ''
                metadata.title = doc_metadata.get('title', '') or ''
                if doc_metadata.get('creationDate'):
                    metadata.created_date = doc_metadata.get('creationDate', '')

            content = '\n\n'.join(text_parts)
            return content, metadata

        finally:
            doc.close()

    async def _process_docx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from Word document"""
        if not self.has_docx:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

        from docx import Document

        doc = Document(file_path)
        metadata = DocumentMetadata()

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        metadata.page_count = len(doc.paragraphs) // 40 + 1  # Estimate

        # Extract tables
        if doc.tables:
            metadata.has_tables = True
            metadata.table_count = len(doc.tables)

            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                text_parts.append('\n' + '\n'.join(table_text))

        # Extract core properties
        core_props = doc.core_properties
        if core_props:
            metadata.author = core_props.author or ''
            metadata.title = core_props.title or ''
            if core_props.created:
                metadata.created_date = str(core_props.created)

        content = '\n\n'.join(text_parts)
        return content, metadata

    async def _process_xlsx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from Excel spreadsheet"""
        if not self.has_openpyxl:
            raise ImportError("openpyxl is required for XLSX processing. Install with: pip install openpyxl")

        import openpyxl

        wb = openpyxl.load_workbook(file_path, data_only=True)
        metadata = DocumentMetadata()

        try:
            text_parts = []
            metadata.sheet_names = wb.sheetnames
            metadata.has_tables = True  # Excel is inherently tabular
            metadata.table_count = len(wb.sheetnames)

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                # Add sheet header
                text_parts.append(f"\n=== Sheet: {sheet_name} ===\n")

                # Extract data
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_text):  # Skip empty rows
                        text_parts.append(' | '.join(row_text))

            content = '\n'.join(text_parts)
            return content, metadata

        finally:
            wb.close()

    async def _process_text(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Process plain text files"""
        metadata = DocumentMetadata()

        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                metadata.language = 'en'
                return content, metadata
            except UnicodeDecodeError:
                continue

        # Fallback to binary read and decode
        with open(file_path, 'rb') as f:
            raw = f.read()
            content = raw.decode('utf-8', errors='replace')

        return content, metadata

    async def process_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process file from bytes (for file uploads).

        Args:
            file_bytes: Raw file content
            filename: Original filename
            content_type: MIME type if known

        Returns:
            ProcessedDocument with extracted content and metadata
        """
        # Save to temp file for processing
        ext = Path(filename).suffix.lower() or '.tmp'

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            result = await self.process_file(tmp_path, filename)
            return result
        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass


class ErrorExplainer:
    """
    Generates human-friendly explanations for document processing errors.
    Uses LLM when available, falls back to pattern-based explanations.

    NOTE: Uses local llama.cpp endpoint only - no external LLM APIs permitted.
    """

    def __init__(self, llm_url: str = "http://localhost:8081", model: str = "general"):
        """
        Initialize ErrorExplainer with local llama.cpp endpoint.

        Args:
            llm_url: Local llama.cpp endpoint (default: 8081 for general model)
            model: Model name for logging purposes
        """
        self.llm_url = llm_url.rstrip('/')
        self.model = model

    def _classify_error(self, error: str) -> str:
        """Classify error type"""
        error_lower = error.lower()

        if 'permission' in error_lower or 'denied' in error_lower:
            return 'permission'
        if 'not found' in error_lower or 'no such file' in error_lower:
            return 'not_found'
        if 'import' in error_lower or 'module' in error_lower:
            return 'dependency'
        if 'decode' in error_lower or 'encoding' in error_lower:
            return 'encoding'
        if 'corrupt' in error_lower or 'invalid' in error_lower:
            return 'corrupt'
        if 'memory' in error_lower or 'size' in error_lower:
            return 'size'

        return 'unknown'

    def get_explanation(self, error: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        Get human-readable explanation for an error.

        Returns:
            Dict with:
                - technical_error: Original error message
                - user_explanation: Human-friendly explanation
                - suggested_actions: List of things to try
                - error_category: Classification of error type
        """
        error_category = self._classify_error(error)
        filename = context.get('filename', 'the file')
        file_type = context.get('file_type', 'document')

        explanations = {
            'permission': {
                'user_explanation': f"I couldn't access {filename} because of permission restrictions.",
                'suggested_actions': [
                    "Check that you have read permissions for the file",
                    "Try copying the file to a different location",
                    "If on a network drive, check your network connection"
                ]
            },
            'not_found': {
                'user_explanation': f"The file {filename} could not be found.",
                'suggested_actions': [
                    "Verify the file exists at the specified location",
                    "Check for typos in the filename",
                    "Try uploading the file again"
                ]
            },
            'dependency': {
                'user_explanation': f"A required library for processing {file_type} files is not installed.",
                'suggested_actions': [
                    f"Install the required Python package (e.g., pip install pymupdf for PDFs)",
                    "Contact your administrator to install the dependency",
                    "Try converting the file to a simpler format like TXT"
                ]
            },
            'encoding': {
                'user_explanation': f"The file {filename} uses an unsupported text encoding.",
                'suggested_actions': [
                    "Save the file with UTF-8 encoding",
                    "Try opening in a text editor and re-saving",
                    "Convert the file to a different format"
                ]
            },
            'corrupt': {
                'user_explanation': f"The file {filename} appears to be damaged or in an invalid format.",
                'suggested_actions': [
                    "Try re-downloading or re-creating the file",
                    "Open the file in its native application to check for errors",
                    "Convert to a different format and try again"
                ]
            },
            'size': {
                'user_explanation': f"The file {filename} is too large to process with available memory.",
                'suggested_actions': [
                    "Try splitting the document into smaller files",
                    "Remove unnecessary images or embedded content",
                    "Contact your administrator for larger file support"
                ]
            },
            'unknown': {
                'user_explanation': f"An unexpected error occurred while processing {filename}.",
                'suggested_actions': [
                    "Try the operation again",
                    "Convert the file to a different format",
                    "Contact support with the error details below"
                ]
            }
        }

        explanation = explanations.get(error_category, explanations['unknown'])

        return {
            'technical_error': error,
            'user_explanation': explanation['user_explanation'],
            'suggested_actions': explanation['suggested_actions'],
            'error_category': error_category
        }

    async def explain_error_with_llm(
        self,
        error: str,
        context: Dict[str, Any]
    ) -> Optional[Dict[str, Any]]:
        """
        Get LLM-powered error explanation (async).
        Falls back to pattern-based if LLM unavailable.
        """
        try:
            import httpx

            prompt = f"""A user tried to upload a document and got this error:

File: {context.get('filename', 'unknown')}
Type: {context.get('file_type', 'unknown')}
Error: {error}

Please explain in 2-3 sentences what went wrong in non-technical terms,
and suggest 2-3 specific actions the user can take to fix this.
Be friendly and helpful."""

            async with httpx.AsyncClient(timeout=10) as client:
                # Use OpenAI-compatible endpoint (llama-cpp-python)
                full_prompt = f"You are a helpful assistant explaining document upload errors to users.\n\n{prompt}"
                response = await client.post(
                    f"{self.llm_url}/v1/completions",
                    json={
                        "prompt": full_prompt,
                        "max_tokens": 200,
                        "temperature": 0.3
                    }
                )

                if response.status_code == 200:
                    data = response.json()
                    # OpenAI format: choices[0].text
                    explanation = data.get('choices', [{}])[0].get('text', '').strip()

                    return {
                        'technical_error': error,
                        'user_explanation': explanation,
                        'suggested_actions': [],  # LLM includes these in explanation
                        'error_category': self._classify_error(error),
                        'llm_generated': True
                    }

        except Exception as e:
            print(f"LLM explanation failed, using pattern-based: {e}")

        # Fall back to pattern-based
        return self.get_explanation(error, context)


# Singleton instance
_document_processor: Optional[DocumentProcessor] = None
_error_explainer: Optional[ErrorExplainer] = None


def get_document_processor() -> DocumentProcessor:
    """Get singleton DocumentProcessor instance"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor


def get_error_explainer(
    llm_url: str = "http://localhost:8081",
    model: str = "general"
) -> ErrorExplainer:
    """Get singleton ErrorExplainer instance"""
    global _error_explainer
    if _error_explainer is None:
        _error_explainer = ErrorExplainer(llm_url, model)
    return _error_explainer
