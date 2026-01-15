"""
EWR Document Agent
==================

Document processing, chunking, embedding, and retrieval agent.
"""

import asyncio
import hashlib
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any, AsyncGenerator

import aiofiles
import numpy as np

from ewr_agent_core import (
    BaseAgent,
    AgentConfig,
    AgentType,
    AgentCapability,
    TaskResult,
    TaskStatus,
)

from .models import (
    Document,
    DocumentType,
    DocumentChunk,
    DocumentMetadata,
    ChunkingConfig,
    ChunkingStrategy,
    ProcessingResult,
    ProcessingStatus,
    SearchResult,
    SearchQuery,
    VectorStoreConfig,
    EmbeddingBatch,
    DocumentBatch,
    ExtractionConfig,
)


# File extension to document type mapping
EXTENSION_MAP = {
    ".pdf": DocumentType.PDF,
    ".docx": DocumentType.DOCX,
    ".doc": DocumentType.DOC,
    ".html": DocumentType.HTML,
    ".htm": DocumentType.HTML,
    ".md": DocumentType.MARKDOWN,
    ".markdown": DocumentType.MARKDOWN,
    ".txt": DocumentType.TEXT,
    ".json": DocumentType.JSON,
    ".xml": DocumentType.XML,
    ".csv": DocumentType.CSV,
    ".py": DocumentType.CODE,
    ".js": DocumentType.CODE,
    ".ts": DocumentType.CODE,
    ".java": DocumentType.CODE,
    ".cs": DocumentType.CODE,
    ".cpp": DocumentType.CODE,
    ".c": DocumentType.CODE,
    ".go": DocumentType.CODE,
    ".rs": DocumentType.CODE,
    ".rb": DocumentType.CODE,
    ".php": DocumentType.CODE,
    ".mp3": DocumentType.AUDIO,
    ".wav": DocumentType.AUDIO,
    ".m4a": DocumentType.AUDIO,
    ".png": DocumentType.IMAGE,
    ".jpg": DocumentType.IMAGE,
    ".jpeg": DocumentType.IMAGE,
}


class DocumentAgent(BaseAgent):
    """
    Document processing agent for ETL pipelines.

    Handles document ingestion, chunking, embedding, and retrieval.
    """

    def __init__(
        self,
        config: Optional[AgentConfig] = None,
        vector_config: Optional[VectorStoreConfig] = None,
        chunking_config: Optional[ChunkingConfig] = None,
        use_local_embeddings: bool = True,
        use_reranking: bool = True,
        embedding_model: Optional[str] = None,
        reranker_model: Optional[str] = None,
    ):
        super().__init__(config=config)

        self.vector_config = vector_config or VectorStoreConfig()
        self.chunking_config = chunking_config or ChunkingConfig()

        # Document cache
        self._document_cache: Dict[str, Document] = {}
        self._chunk_cache: Dict[str, List[DocumentChunk]] = {}

        # Vector store connection (lazy initialized)
        self._vector_store = None

        # Local embedding and reranking services (lazy initialized)
        self._use_local_embeddings = use_local_embeddings
        self._use_reranking = use_reranking
        self._embedding_model = embedding_model
        self._reranker_model = reranker_model
        self._local_embedder = None
        self._reranker = None

    @property
    def agent_type(self) -> AgentType:
        """Agent type identifier."""
        return AgentType.DOCUMENT

    @property
    def capabilities(self) -> List[AgentCapability]:
        """List of capabilities this agent provides."""
        return [
            AgentCapability.DOC_SEARCH,
            AgentCapability.DOC_SUMMARIZE,
            AgentCapability.MULTI_DOC,
        ]

    @property
    def local_embedder(self):
        """Get or initialize local embedding service."""
        if self._local_embedder is None and self._use_local_embeddings:
            try:
                # Try multiple import paths for flexibility
                try:
                    from services.document_embedder import LocalEmbeddingService
                except ImportError:
                    import sys
                    sys.path.insert(0, str(Path(__file__).parents[4] / "python_services"))
                    from services.document_embedder import LocalEmbeddingService
                model = self._embedding_model or LocalEmbeddingService.DEFAULT_MODEL
                self._local_embedder = LocalEmbeddingService(model)
                self.logger.info(f"Local embedder initialized: {model}")
            except ImportError as e:
                self.logger.warning(f"Could not import LocalEmbeddingService: {e}")
        return self._local_embedder

    @property
    def reranker(self):
        """Get or initialize reranking service."""
        if self._reranker is None and self._use_reranking:
            try:
                # Try multiple import paths for flexibility
                try:
                    from services.document_reranker import DocumentRerankerService
                except ImportError:
                    import sys
                    sys.path.insert(0, str(Path(__file__).parents[4] / "python_services"))
                    from services.document_reranker import DocumentRerankerService
                model = self._reranker_model or DocumentRerankerService.DEFAULT_MODEL
                self._reranker = DocumentRerankerService(model)
                self.logger.info(f"Reranker initialized: {model}")
            except ImportError as e:
                self.logger.warning(f"Could not import DocumentRerankerService: {e}")
        return self._reranker

    async def _initialize(self) -> None:
        """Initialize document agent resources."""
        self.logger.info(f"Document agent starting with vector store: {self.vector_config.store_type}")

        # Initialize vector store if configured
        if self.vector_config.store_type == "mongodb":
            await self._init_mongodb()

    async def _agent_stop(self):
        """Cleanup document agent resources."""
        self._document_cache.clear()
        self._chunk_cache.clear()

        # Close MongoDB connection
        if hasattr(self, '_mongo_client') and self._mongo_client:
            self._mongo_client.close()
            self._mongo_client = None
            self._mongo_db = None
            self._vector_store = None
            self.logger.info("MongoDB connection closed")

    async def _init_mongodb(self):
        """Initialize MongoDB vector store."""
        try:
            from motor.motor_asyncio import AsyncIOMotorClient

            connection_string = self.vector_config.connection_string or "mongodb://localhost:27017"
            self._mongo_client = AsyncIOMotorClient(connection_string)
            self._mongo_db = self._mongo_client[self.vector_config.database_name]
            self._vector_store = self._mongo_db[self.vector_config.collection_name]

            # Test connection
            await self._mongo_client.admin.command('ping')
            self.logger.info(f"Connected to MongoDB: {self.vector_config.database_name}/{self.vector_config.collection_name}")

            # Create indexes for efficient search
            await self._vector_store.create_index([("document_id", 1)])
            await self._vector_store.create_index([("chunk_index", 1)])

        except ImportError:
            self.logger.warning("motor not installed. Vector store features disabled.")
        except Exception as e:
            self.logger.error(f"Failed to initialize MongoDB: {e}")

    def _detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension."""
        ext = Path(file_path).suffix.lower()
        return EXTENSION_MAP.get(ext, DocumentType.UNKNOWN)

    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID from file path."""
        return hashlib.sha256(file_path.encode()).hexdigest()[:16]

    async def process_document(
        self,
        file_path: str,
        chunking_config: Optional[ChunkingConfig] = None,
        extraction_config: Optional[ExtractionConfig] = None,
    ) -> ProcessingResult:
        """
        Process a single document through the ETL pipeline.

        Args:
            file_path: Path to the document
            chunking_config: Optional custom chunking configuration
            extraction_config: Optional extraction configuration

        Returns:
            ProcessingResult with status and metadata
        """
        start_time = datetime.utcnow()
        config = chunking_config or self.chunking_config
        extract_config = extraction_config or ExtractionConfig()

        result = ProcessingResult(
            document_id=self._generate_document_id(file_path),
            file_path=file_path,
            status=ProcessingStatus.PARSING,
        )

        try:
            # 1. Parse document
            document = await self._parse_document(file_path, extract_config)
            result.document_type = document.document_type
            result.metadata = document.metadata

            # 2. Chunk document
            result.status = ProcessingStatus.CHUNKING
            chunks = await self._chunk_document(document, config)
            result.total_chunks = len(chunks)

            # 3. Generate embeddings
            result.status = ProcessingStatus.EMBEDDING
            embedded_chunks = await self._embed_chunks(chunks)
            result.chunks_embedded = len([c for c in embedded_chunks if c.embedding])

            # 4. Store in vector database
            result.status = ProcessingStatus.STORING
            await self._store_chunks(embedded_chunks)

            # Cache document
            self._document_cache[document.id] = document
            self._chunk_cache[document.id] = embedded_chunks

            result.status = ProcessingStatus.COMPLETED

        except Exception as e:
            self.logger.error(f"Error processing {file_path}: {e}")
            result.status = ProcessingStatus.FAILED
            result.error = str(e)

        # Calculate processing time
        end_time = datetime.utcnow()
        result.processing_time_ms = int((end_time - start_time).total_seconds() * 1000)

        return result

    async def _parse_document(
        self,
        file_path: str,
        config: ExtractionConfig,
    ) -> Document:
        """Parse a document and extract content."""
        doc_type = self._detect_document_type(file_path)
        doc_id = self._generate_document_id(file_path)

        # Get file info
        path = Path(file_path)
        stat = path.stat()

        metadata = DocumentMetadata(
            title=path.stem,
            file_size=stat.st_size,
            modified_date=datetime.fromtimestamp(stat.st_mtime),
        )

        content = ""

        if doc_type == DocumentType.PDF:
            content, metadata = await self._parse_pdf(file_path, config, metadata)
        elif doc_type == DocumentType.DOCX:
            content, metadata = await self._parse_docx(file_path, config, metadata)
        elif doc_type == DocumentType.HTML:
            content, metadata = await self._parse_html(file_path, config, metadata)
        elif doc_type == DocumentType.MARKDOWN:
            content = await self._parse_markdown(file_path)
        elif doc_type in (DocumentType.TEXT, DocumentType.CODE):
            content = await self._parse_text(file_path)
        elif doc_type == DocumentType.JSON:
            content = await self._parse_json(file_path)
        else:
            # Try as text
            content = await self._parse_text(file_path)

        # Update word count
        metadata.word_count = len(content.split())

        return Document(
            id=doc_id,
            file_path=file_path,
            document_type=doc_type,
            content=content,
            metadata=metadata,
        )

    async def _parse_pdf(
        self,
        file_path: str,
        config: ExtractionConfig,
        metadata: DocumentMetadata,
    ) -> tuple[str, DocumentMetadata]:
        """Parse PDF document."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            metadata.page_count = len(reader.pages)

            # Extract metadata
            if reader.metadata:
                metadata.title = reader.metadata.get("/Title") or metadata.title
                metadata.author = reader.metadata.get("/Author")

            # Extract text from all pages
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {i+1}]\n{page_text}")

            return "\n\n".join(text_parts), metadata

        except ImportError:
            self.logger.warning("pypdf not installed. Cannot parse PDF.")
            return "", metadata
        except Exception as e:
            self.logger.error(f"Error parsing PDF {file_path}: {e}")
            return "", metadata

    async def _parse_docx(
        self,
        file_path: str,
        config: ExtractionConfig,
        metadata: DocumentMetadata,
    ) -> tuple[str, DocumentMetadata]:
        """Parse DOCX document."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)

            # Extract core properties
            if doc.core_properties:
                metadata.title = doc.core_properties.title or metadata.title
                metadata.author = doc.core_properties.author
                metadata.created_date = doc.core_properties.created
                metadata.modified_date = doc.core_properties.modified

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Extract tables if configured
            if config.extract_tables:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        table_text.append(row_text)
                    if table_text:
                        paragraphs.append("\n[Table]\n" + "\n".join(table_text))

            return "\n\n".join(paragraphs), metadata

        except ImportError:
            self.logger.warning("python-docx not installed. Cannot parse DOCX.")
            return "", metadata
        except Exception as e:
            self.logger.error(f"Error parsing DOCX {file_path}: {e}")
            return "", metadata

    async def _parse_html(
        self,
        file_path: str,
        config: ExtractionConfig,
        metadata: DocumentMetadata,
    ) -> tuple[str, DocumentMetadata]:
        """Parse HTML document."""
        try:
            from bs4 import BeautifulSoup

            async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                html_content = await f.read()

            soup = BeautifulSoup(html_content, "lxml")

            # Extract title
            title_tag = soup.find("title")
            if title_tag:
                metadata.title = title_tag.get_text().strip()

            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()

            # Extract text
            text = soup.get_text(separator="\n")

            # Clean up whitespace
            lines = [line.strip() for line in text.splitlines()]
            text = "\n".join(line for line in lines if line)

            # Extract links if configured
            if config.extract_links:
                links = []
                for a in soup.find_all("a", href=True):
                    link_text = a.get_text().strip()
                    href = a["href"]
                    if link_text and href and not href.startswith("#"):
                        links.append(f"[{link_text}]({href})")
                if links:
                    text += "\n\n[Links]\n" + "\n".join(links[:20])

            return text, metadata

        except ImportError:
            self.logger.warning("beautifulsoup4 not installed. Cannot parse HTML.")
            return "", metadata
        except Exception as e:
            self.logger.error(f"Error parsing HTML {file_path}: {e}")
            return "", metadata

    async def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown document."""
        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return await f.read()

    async def _parse_text(self, file_path: str) -> str:
        """Parse plain text document."""
        async with aiofiles.open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return await f.read()

    async def _parse_json(self, file_path: str) -> str:
        """Parse JSON document."""
        import json

        async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
            content = await f.read()

        # Pretty print for better chunking
        data = json.loads(content)
        return json.dumps(data, indent=2)

    async def _chunk_document(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk a document based on configuration."""
        if config.strategy == ChunkingStrategy.FIXED_SIZE:
            return self._chunk_fixed_size(document, config)
        elif config.strategy == ChunkingStrategy.SENTENCE:
            return self._chunk_by_sentences(document, config)
        elif config.strategy == ChunkingStrategy.PARAGRAPH:
            return self._chunk_by_paragraphs(document, config)
        elif config.strategy == ChunkingStrategy.RECURSIVE:
            return self._chunk_recursive(document, config)
        elif config.strategy == ChunkingStrategy.MARKDOWN:
            return self._chunk_markdown(document, config)
        elif config.strategy == ChunkingStrategy.CODE:
            return self._chunk_code(document, config)
        else:
            return self._chunk_recursive(document, config)

    def _chunk_fixed_size(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk by fixed character size."""
        chunks = []
        content = document.content
        step = config.chunk_size - config.chunk_overlap

        for i in range(0, len(content), step):
            chunk_content = content[i:i + config.chunk_size]

            if len(chunk_content) < config.min_chunk_size:
                # Append to last chunk if too small
                if chunks:
                    chunks[-1].content += chunk_content
                    chunks[-1].end_char = i + len(chunk_content)
                continue

            chunks.append(DocumentChunk(
                id=f"{document.id}_{len(chunks)}",
                document_id=document.id,
                content=chunk_content,
                chunk_index=len(chunks),
                start_char=i,
                end_char=i + len(chunk_content),
                metadata={"source": document.file_path},
            ))

        return chunks

    def _chunk_by_sentences(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk by sentences."""
        # Simple sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', document.content)

        chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence_len = len(sentence)

            if current_length + sentence_len > config.chunk_size and current_chunk:
                # Save current chunk
                chunk_content = " ".join(current_chunk)
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    metadata={"source": document.file_path},
                ))
                current_chunk = []
                current_length = 0

            current_chunk.append(sentence)
            current_length += sentence_len + 1

        # Don't forget last chunk
        if current_chunk:
            chunk_content = " ".join(current_chunk)
            if len(chunk_content) >= config.min_chunk_size:
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    metadata={"source": document.file_path},
                ))
            elif chunks:
                chunks[-1].content += " " + chunk_content

        return chunks

    def _chunk_by_paragraphs(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk by paragraphs."""
        paragraphs = document.content.split("\n\n")

        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_len = len(para)

            if current_length + para_len > config.chunk_size and current_chunk:
                chunk_content = "\n\n".join(current_chunk)
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    metadata={"source": document.file_path},
                ))
                current_chunk = []
                current_length = 0

            current_chunk.append(para)
            current_length += para_len + 2

        if current_chunk:
            chunk_content = "\n\n".join(current_chunk)
            if len(chunk_content) >= config.min_chunk_size:
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=len(chunks),
                    metadata={"source": document.file_path},
                ))
            elif chunks:
                chunks[-1].content += "\n\n" + chunk_content

        return chunks

    def _chunk_recursive(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Recursive text splitter - tries larger separators first."""
        chunks = []

        def split_recursive(text: str, separators: List[str]) -> List[str]:
            if not separators:
                return [text]

            separator = separators[0]
            remaining = separators[1:]

            parts = text.split(separator) if separator else list(text)
            result = []
            current = []
            current_len = 0

            for part in parts:
                part_len = len(part)

                if current_len + part_len + len(separator) > config.chunk_size:
                    if current:
                        combined = separator.join(current)
                        if len(combined) > config.max_chunk_size and remaining:
                            result.extend(split_recursive(combined, remaining))
                        else:
                            result.append(combined)
                    current = [part]
                    current_len = part_len
                else:
                    current.append(part)
                    current_len += part_len + len(separator)

            if current:
                combined = separator.join(current)
                if len(combined) > config.max_chunk_size and remaining:
                    result.extend(split_recursive(combined, remaining))
                else:
                    result.append(combined)

            return result

        text_chunks = split_recursive(document.content, config.separators)

        for i, chunk_content in enumerate(text_chunks):
            if len(chunk_content) >= config.min_chunk_size:
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{i}",
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=i,
                    metadata={"source": document.file_path},
                ))

        return chunks

    def _chunk_markdown(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk markdown by headers."""
        chunks = []
        content = document.content

        # Split by headers
        header_pattern = r'^(#{1,6})\s+(.+)$'
        sections = re.split(r'(?m)^(?=#{1,6}\s)', content)

        current_header = ""

        for section in sections:
            section = section.strip()
            if not section:
                continue

            # Extract header if present
            header_match = re.match(header_pattern, section, re.MULTILINE)
            if header_match:
                current_header = header_match.group(2)

            if len(section) > config.max_chunk_size:
                # Split large sections
                sub_chunks = self._chunk_by_paragraphs(
                    Document(id=document.id, file_path=document.file_path, content=section),
                    config
                )
                for sub in sub_chunks:
                    sub.section = current_header
                    sub.id = f"{document.id}_{len(chunks)}"
                    sub.chunk_index = len(chunks)
                    chunks.append(sub)
            elif len(section) >= config.min_chunk_size:
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=section,
                    chunk_index=len(chunks),
                    section=current_header,
                    metadata={"source": document.file_path},
                ))

        return chunks

    def _chunk_code(
        self,
        document: Document,
        config: ChunkingConfig,
    ) -> List[DocumentChunk]:
        """Chunk code by functions/classes."""
        chunks = []
        content = document.content

        # Simple function/class detection (works for Python, JS, etc.)
        patterns = [
            r'(?m)^(class\s+\w+.*?(?=\nclass\s|\n(?:def|function)\s|$))',  # Classes
            r'(?m)^((?:async\s+)?(?:def|function)\s+\w+.*?(?=\n(?:def|function|class)\s|$))',  # Functions
        ]

        # Find all code blocks
        blocks = []
        used_ranges = set()

        for pattern in patterns:
            for match in re.finditer(pattern, content, re.DOTALL):
                start, end = match.span()
                # Avoid overlapping
                if any(start < e and end > s for s, e in used_ranges):
                    continue
                blocks.append((start, end, match.group(0)))
                used_ranges.add((start, end))

        # Sort by position
        blocks.sort(key=lambda x: x[0])

        # Create chunks
        for i, (start, end, block_content) in enumerate(blocks):
            if len(block_content) > config.max_chunk_size:
                # Split large blocks
                for j in range(0, len(block_content), config.chunk_size):
                    sub_content = block_content[j:j + config.chunk_size]
                    if len(sub_content) >= config.min_chunk_size:
                        chunks.append(DocumentChunk(
                            id=f"{document.id}_{len(chunks)}",
                            document_id=document.id,
                            content=sub_content,
                            chunk_index=len(chunks),
                            start_char=start + j,
                            end_char=start + j + len(sub_content),
                            metadata={"source": document.file_path, "type": "code"},
                        ))
            elif len(block_content) >= config.min_chunk_size:
                chunks.append(DocumentChunk(
                    id=f"{document.id}_{len(chunks)}",
                    document_id=document.id,
                    content=block_content,
                    chunk_index=len(chunks),
                    start_char=start,
                    end_char=end,
                    metadata={"source": document.file_path, "type": "code"},
                ))

        # If no blocks found, fall back to recursive
        if not chunks:
            return self._chunk_recursive(document, config)

        return chunks

    async def _embed_chunks(
        self,
        chunks: List[DocumentChunk],
    ) -> List[DocumentChunk]:
        """Generate embeddings for chunks using local embedder or LLM backend."""
        if not chunks:
            return chunks

        # Try local embedder first (faster, no API calls)
        if self.local_embedder:
            try:
                texts = [chunk.content for chunk in chunks]
                embeddings = await self.local_embedder.embed_batch(texts)
                for chunk, embedding in zip(chunks, embeddings):
                    chunk.embedding = embedding
                    chunk.embedding_model = self.local_embedder.model_name
                self.logger.debug(f"Embedded {len(chunks)} chunks with local embedder")
                return chunks
            except Exception as e:
                self.logger.warning(f"Local embedder failed, falling back to LLM: {e}")

        # Fall back to LLM backend
        if not self.llm:
            self.logger.warning("No LLM backend configured for embeddings")
            return chunks

        for chunk in chunks:
            try:
                embedding = await self.llm.embed(chunk.content)
                chunk.embedding = embedding
                chunk.embedding_model = self.llm.model
            except Exception as e:
                self.logger.error(f"Error generating embedding for chunk {chunk.id}: {e}")

        return chunks

    async def _store_chunks(self, chunks: List[DocumentChunk]):
        """Store chunks in MongoDB."""
        if self._vector_store is None:
            self.logger.debug("No vector store configured, skipping storage")
            return

        try:
            # Prepare documents for MongoDB
            documents = []
            for chunk in chunks:
                if chunk.embedding:
                    documents.append({
                        "_id": chunk.id,
                        "document_id": chunk.document_id,
                        "content": chunk.content,
                        "chunk_index": chunk.chunk_index,
                        "vector": chunk.embedding,
                        "metadata": chunk.metadata,
                        "section": chunk.section,
                        "start_char": chunk.start_char,
                        "end_char": chunk.end_char,
                        "embedding_model": chunk.embedding_model,
                        "created_at": chunk.created_at.isoformat(),
                    })

            if documents:
                # Delete existing chunks for this document first
                if documents:
                    doc_id = documents[0]["document_id"]
                    await self._vector_store.delete_many({"document_id": doc_id})

                # Insert new chunks
                await self._vector_store.insert_many(documents)
                self.logger.info(f"Stored {len(documents)} chunks in MongoDB")

        except Exception as e:
            self.logger.error(f"Error storing chunks in MongoDB: {e}")

    async def search(
        self,
        query: str,
        top_k: int = 5,
        min_score: float = 0.0,
        filter_document_ids: Optional[List[str]] = None,
        use_reranking: Optional[bool] = None,
    ) -> List[SearchResult]:
        """
        Search for relevant document chunks.

        Args:
            query: Search query text
            top_k: Number of results to return
            min_score: Minimum similarity score
            filter_document_ids: Optional list of document IDs to filter
            use_reranking: Override instance reranking setting (None = use default)

        Returns:
            List of SearchResult with matching chunks
        """
        results = []
        should_rerank = use_reranking if use_reranking is not None else self._use_reranking

        # Generate query embedding - try local embedder first
        query_embedding = None
        if self.local_embedder:
            try:
                query_embedding = await self.local_embedder.embed(query)
                self.logger.debug("Query embedded with local embedder")
            except Exception as e:
                self.logger.warning(f"Local embedder failed for query: {e}")

        # Fall back to LLM backend
        if query_embedding is None:
            if not self.llm:
                self.logger.warning("No embedding backend available")
                return results
            try:
                query_embedding = await self.llm.embed(query)
            except Exception as e:
                self.logger.error(f"Failed to embed query: {e}")
                return results

        # Fetch more candidates if reranking (cross-encoder will refine)
        candidate_k = top_k * 10 if should_rerank and self.reranker else top_k

        try:
            if self._vector_store is not None:
                # Search MongoDB using in-memory cosine similarity
                # Build filter
                query_filter = {}
                if filter_document_ids:
                    query_filter["document_id"] = {"$in": filter_document_ids}

                # Fetch all chunks with embeddings
                cursor = self._vector_store.find(
                    {**query_filter, "vector": {"$exists": True}},
                    {"_id": 1, "document_id": 1, "content": 1, "chunk_index": 1, "vector": 1, "metadata": 1}
                )

                search_results = []
                async for doc in cursor:
                    if doc.get("vector"):
                        score = self._cosine_similarity(query_embedding, doc["vector"])
                        if score >= min_score:
                            search_results.append((score, doc))

                # Sort by score and limit to candidates
                search_results.sort(key=lambda x: x[0], reverse=True)
                search_results = search_results[:candidate_k]

                for score, row in search_results:
                    chunk = DocumentChunk(
                        id=row["_id"],
                        document_id=row["document_id"],
                        content=row["content"],
                        chunk_index=row["chunk_index"],
                    )
                    metadata = row.get("metadata", {})
                    doc_path = metadata.get("source", "") if isinstance(metadata, dict) else str(metadata)
                    results.append(SearchResult(
                        chunk=chunk,
                        score=score,
                        document_path=doc_path,
                    ))
            else:
                # Fall back to cached chunks
                for doc_id, chunks in self._chunk_cache.items():
                    if filter_document_ids and doc_id not in filter_document_ids:
                        continue

                    for chunk in chunks:
                        if chunk.embedding:
                            # Simple cosine similarity
                            score = self._cosine_similarity(query_embedding, chunk.embedding)
                            if score >= min_score:
                                doc = self._document_cache.get(doc_id)
                                results.append(SearchResult(
                                    chunk=chunk,
                                    score=score,
                                    document_path=doc.file_path if doc else "",
                                ))

                # Sort by score and limit to candidates
                results.sort(key=lambda x: x.score, reverse=True)
                results = results[:candidate_k]

            # Apply cross-encoder re-ranking if enabled
            if should_rerank and self.reranker and len(results) > 0:
                results = await self._rerank_results(query, results, top_k)
            else:
                # Just limit to top_k
                results = results[:top_k]

        except Exception as e:
            self.logger.error(f"Error searching: {e}")

        return results

    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors using numpy."""
        if len(vec1) != len(vec2):
            return 0.0

        v1 = np.array(vec1)
        v2 = np.array(vec2)

        norm1 = np.linalg.norm(v1)
        norm2 = np.linalg.norm(v2)

        if norm1 == 0 or norm2 == 0:
            return 0.0

        return float(np.dot(v1, v2) / (norm1 * norm2))

    async def _rerank_results(
        self,
        query: str,
        results: List[SearchResult],
        top_k: int,
    ) -> List[SearchResult]:
        """
        Re-rank search results using cross-encoder model.

        Args:
            query: Original search query
            results: Initial bi-encoder results
            top_k: Number of final results to return

        Returns:
            Re-ranked list of SearchResult
        """
        if not self.reranker or not results:
            return results[:top_k]

        try:
            # Extract content for re-ranking
            documents = [r.chunk.content for r in results]

            # Get re-ranked indices and scores
            ranked = await self.reranker.rerank(query, documents, top_k)

            # Build re-ranked results with cross-encoder scores
            reranked_results = []
            for idx, ce_score in ranked:
                result = results[idx]
                # Update score with cross-encoder score (normalized to 0-1)
                # Cross-encoder scores can be negative, so we normalize
                normalized_score = 1.0 / (1.0 + np.exp(-ce_score))  # sigmoid
                reranked_results.append(SearchResult(
                    chunk=result.chunk,
                    score=float(normalized_score),
                    document_path=result.document_path,
                ))

            self.logger.debug(f"Re-ranked {len(results)} results to top {len(reranked_results)}")
            return reranked_results

        except Exception as e:
            self.logger.error(f"Re-ranking failed, using bi-encoder scores: {e}")
            return results[:top_k]

    async def process_folder(
        self,
        folder_path: str,
        pattern: str = "*.*",
        recursive: bool = True,
        chunking_config: Optional[ChunkingConfig] = None,
    ) -> List[ProcessingResult]:
        """
        Process all documents in a folder.

        Args:
            folder_path: Path to folder
            pattern: Glob pattern for files
            recursive: Search recursively
            chunking_config: Optional chunking configuration

        Returns:
            List of ProcessingResult for each document
        """
        results = []
        folder = Path(folder_path)

        if not folder.exists():
            self.logger.error(f"Folder not found: {folder_path}")
            return results

        # Find files
        if recursive:
            files = list(folder.rglob(pattern))
        else:
            files = list(folder.glob(pattern))

        # Filter to supported types
        supported_files = [
            f for f in files
            if f.suffix.lower() in EXTENSION_MAP
        ]

        self.logger.info(f"Processing {len(supported_files)} files from {folder_path}")

        for file_path in supported_files:
            try:
                result = await self.process_document(
                    str(file_path),
                    chunking_config=chunking_config,
                )
                results.append(result)
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {e}")
                results.append(ProcessingResult(
                    document_id=self._generate_document_id(str(file_path)),
                    file_path=str(file_path),
                    status=ProcessingStatus.FAILED,
                    error=str(e),
                ))

        return results

    async def get_document(self, document_id: str) -> Optional[Document]:
        """Get a cached document by ID."""
        return self._document_cache.get(document_id)

    async def get_chunks(self, document_id: str) -> List[DocumentChunk]:
        """Get chunks for a document."""
        return self._chunk_cache.get(document_id, [])

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks."""
        if document_id in self._document_cache:
            del self._document_cache[document_id]

        if document_id in self._chunk_cache:
            del self._chunk_cache[document_id]

        # Delete from MongoDB
        if self._vector_store is not None:
            try:
                result = await self._vector_store.delete_many({"document_id": document_id})
                self.logger.info(f"Deleted {result.deleted_count} chunks for document {document_id}")
            except Exception as e:
                self.logger.error(f"Error deleting from MongoDB: {e}")
                return False

        return True

    async def handle_task(self, task: Dict[str, Any]) -> TaskResult:
        """
        Handle an incoming task from the message broker or direct call.

        Supported task types:
        - process_document: Process a document file
        - search: Search for relevant chunks
        - delete: Delete a document

        Args:
            task: Task dictionary with 'type' and parameters

        Returns:
            TaskResult with status and result data
        """
        task_id = task.get("id", "unknown")
        task_type = task.get("type", "unknown")

        try:
            if task_type == "process_document":
                file_path = task.get("file_path")
                if not file_path:
                    return TaskResult(
                        task_id=task_id,
                        status=TaskStatus.FAILED,
                        error="file_path is required for process_document task"
                    )
                result = await self.process_document(file_path)
                return TaskResult(
                    task_id=task_id,
                    status=TaskStatus.COMPLETED if result.status == ProcessingStatus.COMPLETED else TaskStatus.FAILED,
                    result=result.model_dump() if hasattr(result, 'model_dump') else vars(result),
                    error=result.error
                )

            elif task_type == "search":
                query = task.get("query")
                if not query:
                    return TaskResult(
                        task_id=task_id,
                        status=TaskStatus.FAILED,
                        error="query is required for search task"
                    )
                results = await self.search(
                    query=query,
                    top_k=task.get("top_k", 5),
                    min_score=task.get("min_score", 0.0),
                    filter_document_ids=task.get("filter_document_ids")
                )
                return TaskResult(
                    task_id=task_id,
                    status=TaskStatus.COMPLETED,
                    result={"results": [r.model_dump() if hasattr(r, 'model_dump') else vars(r) for r in results]}
                )

            elif task_type == "delete":
                document_id = task.get("document_id")
                if not document_id:
                    return TaskResult(
                        task_id=task_id,
                        status=TaskStatus.FAILED,
                        error="document_id is required for delete task"
                    )
                success = await self.delete_document(document_id)
                return TaskResult(
                    task_id=task_id,
                    status=TaskStatus.COMPLETED if success else TaskStatus.FAILED,
                    result={"deleted": success, "document_id": document_id}
                )

            else:
                return TaskResult(
                    task_id=task_id,
                    status=TaskStatus.FAILED,
                    error=f"Unknown task type: {task_type}"
                )

        except Exception as e:
            self.logger.error(f"Error handling task {task_id}: {e}")
            return TaskResult(
                task_id=task_id,
                status=TaskStatus.FAILED,
                error=str(e)
            )

    async def get_stats(self) -> Dict[str, Any]:
        """Get document processing statistics."""
        return {
            "cached_documents": len(self._document_cache),
            "cached_chunks": sum(len(c) for c in self._chunk_cache.values()),
            "vector_store": self.vector_config.store_type,
            "vector_store_connected": self._vector_store is not None,
            "chunking_strategy": self.chunking_config.strategy.value,
            "chunk_size": self.chunking_config.chunk_size,
        }
